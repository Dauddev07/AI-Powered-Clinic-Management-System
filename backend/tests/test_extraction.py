import io

import docx

from app.rag.extraction import _extract_docx


def _docx_bytes_with_table(paragraph_before, headers, rows, paragraph_after=None):
    document = docx.Document()
    document.add_paragraph(paragraph_before)

    table = document.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row_values in rows:
        row_cells = table.add_row().cells
        for cell, value in zip(row_cells, row_values):
            cell.text = value

    if paragraph_after:
        document.add_paragraph(paragraph_after)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_paragraph_only_docx_extracts_as_before():
    document = docx.Document()
    document.add_paragraph("The clinic is open Monday to Friday from 9am to 5pm.")
    buf = io.BytesIO()
    document.save(buf)

    text = _extract_docx(buf.getvalue())

    assert "The clinic is open Monday to Friday from 9am to 5pm." in text


def test_table_content_is_extracted_not_silently_dropped():
    # Reproduces the original bug's shape: doctor schedules stored in a Word table
    # were entirely absent from extracted text because _extract_docx only walked
    # document.paragraphs and never touched document.tables.
    file_bytes = _docx_bytes_with_table(
        paragraph_before="Cardiology Department",
        headers=["Doctor", "Shift"],
        rows=[
            ["Dr. Ayesha Khan", "Mon/Wed/Fri 9am-1pm"],
            ["Dr. Bilal Ahmed", "Tue/Thu 2pm-6pm"],
        ],
    )

    text = _extract_docx(file_bytes)

    assert "Cardiology Department" in text
    assert "Dr. Ayesha Khan" in text
    assert "Mon/Wed/Fri 9am-1pm" in text
    assert "Dr. Bilal Ahmed" in text
    assert "Tue/Thu 2pm-6pm" in text


def test_table_rows_keep_cells_associated_with_each_other():
    # Not just present *somewhere* in the text — each doctor's name must stay on the
    # same line as their own shift, not get flattened into an unassociated word soup
    # where "Dr. Ayesha Khan" could be misread as paired with the wrong shift.
    file_bytes = _docx_bytes_with_table(
        paragraph_before="Cardiology Department",
        headers=["Doctor", "Shift"],
        rows=[
            ["Dr. Ayesha Khan", "Mon/Wed/Fri 9am-1pm"],
            ["Dr. Bilal Ahmed", "Tue/Thu 2pm-6pm"],
        ],
    )

    text = _extract_docx(file_bytes)
    lines = text.splitlines()

    ayesha_line = next(line for line in lines if "Ayesha" in line)
    bilal_line = next(line for line in lines if "Bilal" in line)

    assert "Mon/Wed/Fri 9am-1pm" in ayesha_line
    assert "Tue/Thu 2pm-6pm" not in ayesha_line
    assert "Tue/Thu 2pm-6pm" in bilal_line


def test_table_content_is_interleaved_with_surrounding_paragraphs_in_document_order():
    file_bytes = _docx_bytes_with_table(
        paragraph_before="Cardiology Department",
        headers=["Doctor", "Shift"],
        rows=[["Dr. Ayesha Khan", "Mon/Wed/Fri 9am-1pm"]],
        paragraph_after="Neurology Department",
    )

    text = _extract_docx(file_bytes)

    before_idx = text.index("Cardiology Department")
    table_idx = text.index("Dr. Ayesha Khan")
    after_idx = text.index("Neurology Department")

    assert before_idx < table_idx < after_idx
