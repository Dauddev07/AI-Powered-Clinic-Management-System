import io
import uuid

import docx
import pytest

from app.models.clinic import Clinic
from app.rag.chroma_client import get_chroma_client, get_hospital_info_collection, hospital_info_collection_name
from app.rag.embeddings import LocalEmbeddings, embed_query
from app.rag.preprocessing import preprocess_query
from app.rag.retrieval import FALLBACK_MESSAGE, _best_dense_score, _build_hybrid_retriever, expand_query, retrieve
from app.services.kb_documents import ingest_document


@pytest.fixture
def clinic(db):
    slug = f"test-{uuid.uuid4().hex[:8]}"
    c = Clinic(name="Test Clinic", slug=slug)
    db.add(c)
    db.flush()
    yield c
    # Chroma collections live outside the Postgres transaction the `db` fixture rolls
    # back, so they must be cleaned up explicitly or every test run leaks collections
    # into the shared dockerized ChromaDB instance.
    client = get_chroma_client()
    try:
        client.delete_collection(hospital_info_collection_name(slug))
    except Exception:
        pass


def _ingest(db, clinic, filename, text):
    # ingest_document extracts from real PDF/DOCX bytes; retrieval tests only care
    # about what lands in Chroma, so bypass extraction and call the shared
    # chunk/clean/embed/upsert steps directly against the clinic-scoped collection.
    from app.rag.chunking import chunk_text
    from app.rag.embeddings import embed_texts
    from app.rag.preprocessing import clean_text
    from app.models.kb_document import KBDocument

    collection = get_hospital_info_collection(db, clinic.id)
    chunks = chunk_text(text)
    cleaned = [clean_text(c) for c in chunks]
    embeddings = embed_texts(cleaned)

    document = KBDocument(
        clinic_id=clinic.id,
        source_type="pdf",
        source_ref_id=None,
        title=filename,
        content=text,
        chroma_collection=collection.name,
        chunk_count=len(cleaned),
    )
    db.add(document)
    db.flush()

    ids = [f"{document.id}_{i}" for i in range(len(cleaned))]
    metadatas = [{"kb_document_id": str(document.id), "clinic_id": str(clinic.id), "filename": filename, "chunk_index": i} for i in range(len(cleaned))]
    collection.upsert(ids=ids, embeddings=embeddings, documents=cleaned, metadatas=metadatas)
    return document


def _docx_bytes(text: str) -> bytes:
    document = docx.Document()
    document.add_paragraph(text)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# --- ingestion lands in the hospital_info collection --------------------------------


def test_upload_lands_in_hospital_info_collection(db, clinic):
    result = ingest_document(
        db, clinic.id, "clinic-timings.docx",
        _docx_bytes("The clinic is open Monday to Friday from 9am to 5pm."),
    )

    assert result.document.chroma_collection == hospital_info_collection_name(clinic.slug)

    hospital_collection = get_hospital_info_collection(db, clinic.id)
    assert hospital_collection.count() > 0


# --- grounded reply vs. fallback -----------------------------------------------------


def test_timings_question_retrieves_matching_chunk(db, clinic):
    _ingest(db, clinic, "timings.pdf", "The clinic opening hours are 9am to 5pm, Monday through Friday.")

    result = retrieve(db, clinic.id, "What time does the clinic open?")

    assert result.matched is True
    assert any("opening hours" in c or "9am" in c for c in result.chunks)


def test_out_of_scope_query_returns_fixed_fallback_verbatim(db, clinic):
    _ingest(db, clinic, "timings.pdf", "The clinic opening hours are 9am to 5pm, Monday through Friday.")

    result = retrieve(db, clinic.id, "What is the capital of France and how do transistors work?")

    assert result.matched is False
    assert result.fallback_message == FALLBACK_MESSAGE
    assert result.chunks == []


# --- BM25 catches an exact name dense search alone can miss -------------------------


def test_bm25_surfaces_exact_code_that_dense_only_search_ranks_below_a_paraphrase(db, clinic):
    # An opaque tracking code shares no real semantic content with the query beyond
    # the literal token, while a paraphrased distractor shares heavy topical/lexical
    # overlap ("consultation", "Thursday afternoon") without the code. A dense-only
    # search drifts toward the fluent paraphrase; BM25's exact-token match still nails
    # the coded document — proving BM25 (not the embedding model) is what surfaces it.
    code = "qzx8841krn"
    coded_doc_text = f"Internal tracking reference {code}. Do not distribute externally."
    distractor_text = "Consultations with Dr. Amara Osei are available every Thursday afternoon for follow-up visits."
    _ingest(db, clinic, "coded-doc.pdf", coded_doc_text)
    _ingest(db, clinic, "distractor.pdf", distractor_text)

    collection = get_hospital_info_collection(db, clinic.id)
    query = f"Is there a consultation available with reference {code} on Thursday afternoon?"

    preprocessed = preprocess_query(query)

    from langchain_community.vectorstores import Chroma
    vectorstore = Chroma(client=get_chroma_client(), collection_name=collection.name, embedding_function=LocalEmbeddings())
    dense_only_top = vectorstore.as_retriever(search_kwargs={"k": 1}).invoke(preprocessed)

    ensemble = _build_hybrid_retriever(collection)
    ensemble_results = ensemble.invoke(preprocessed)

    dense_only_hit_code = any(code in d.page_content.lower() for d in dense_only_top)
    ensemble_hit_code = any(code in d.page_content.lower() for d in ensemble_results)

    assert ensemble_hit_code, "BM25 should surface the exact-code document via the ensemble retriever"
    assert not dense_only_hit_code, "test is only meaningful if dense-only search ranks the paraphrased distractor above the coded document"


# --- clinic isolation ----------------------------------------------------------------


def test_clinic_a_retrieval_returns_nothing_from_clinic_b(db, clinic):
    other_slug = f"test-{uuid.uuid4().hex[:8]}"
    other = Clinic(name="Other Clinic", slug=other_slug)
    db.add(other)
    db.flush()
    try:
        _ingest(db, other, "secret.pdf", "Clinic B's confidential internal fee schedule and doctor roster details.")
        _ingest(db, clinic, "timings.pdf", "The clinic opening hours are 9am to 5pm, Monday through Friday.")

        result = retrieve(db, clinic.id, "Tell me about the confidential internal fee schedule and doctor roster.")

        assert all("confidential" not in c for c in result.chunks)
        assert all("Clinic B" not in c for c in result.chunks)
    finally:
        client = get_chroma_client()
        try:
            client.delete_collection(hospital_info_collection_name(other_slug))
        except Exception:
            pass


# --- query expansion ------------------------------------------------------------


def test_query_expansion_pushes_a_below_threshold_query_above_it(db, clinic):
    # The KB chunk is phrased formally ("reservation"); the patient's query uses a
    # colloquial term ("book") instead. Without expansion this genuinely scores below
    # RETRIEVAL_SIMILARITY_THRESHOLD — expansion is what closes the gap.
    _ingest(
        db, clinic, "booking.pdf",
        "To make a reservation, log into the patient portal, select your preferred department, "
        "choose an available slot, and confirm.",
    )

    query = "how do I book a visit"
    preprocessed = preprocess_query(query)
    collection = get_hospital_info_collection(db, clinic.id)
    unexpanded_score = _best_dense_score(collection, embed_query(preprocessed))

    result = retrieve(db, clinic.id, query)

    assert result.best_score >= unexpanded_score
    assert result.matched is True
    assert any("reservation" in c or "portal" in c for c in result.chunks)


def test_irrelevant_query_still_falls_back_even_when_expansion_fires(db, clinic):
    # "price" is both a hospital-info-relevant term and an expansion dictionary key
    # (-> "fee", "cost", "charge"), so expansion actively fires here — the assertion is
    # that adding those terms still isn't enough to make an unrelated question about
    # laptop prices match the clinic's own fee/hours content.
    _ingest(db, clinic, "timings.pdf", "The clinic opening hours are 9am to 5pm, Monday through Friday.")

    query = "Can I get a good price on a new laptop?"
    assert expand_query(preprocess_query(query)) != preprocess_query(query), "test is only meaningful if expansion actually fires for this query"

    result = retrieve(db, clinic.id, query)

    assert result.matched is False
    assert result.fallback_message == FALLBACK_MESSAGE
    assert result.chunks == []
