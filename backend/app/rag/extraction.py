"""Raw text extraction for supported knowledge-base document formats."""
import io

import docx
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


class UnsupportedDocumentType(Exception):
    pass


def extract_text(filename: str, file_bytes: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    raise UnsupportedDocumentType(f"Unsupported file type: {filename} (only .pdf and .docx are accepted)")


def _extract_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _iter_block_items(parent):
    """Yields Paragraph/Table objects from a Document (or table Cell) in document
    order. python-docx's own `.paragraphs`/`.tables` accessors each only walk one
    element type in isolation from the other — this walks the underlying XML body
    directly so nothing between/around a table is skipped or reordered."""
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError(f"unsupported parent type for block iteration: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _table_to_text(table: Table) -> str:
    """Flattens a table into plain text, one line per row, cells joined with " | "
    so column relationships (e.g. doctor name | department | shift) survive as
    readable text rather than being lost along with the table structure. Cells
    containing a nested table are recursed into so nothing inside them is dropped."""
    lines = []
    for row in table.rows:
        cell_texts = []
        for cell in row.cells:
            nested = "\n".join(_table_to_text(t) for t in cell.tables)
            cell_text = "\n".join(filter(None, [cell.text.strip(), nested]))
            if cell_text:
                cell_texts.append(cell_text)
        if cell_texts:
            lines.append(" | ".join(cell_texts))
    return "\n".join(lines)


def _extract_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            if block.text.strip():
                parts.append(block.text)
        elif isinstance(block, Table):
            table_text = _table_to_text(block)
            if table_text:
                parts.append(table_text)
    # Blank-line joined (not single "\n") so downstream chunking's paragraph-boundary
    # separator ("\n\n") sees each paragraph/table as its own preferred split point,
    # instead of greedily packing unrelated paragraphs into one chunk together.
    return "\n\n".join(parts)
